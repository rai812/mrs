# coding=utf-8
from django.conf import settings
import os
from django.http import HttpResponse

from docx import Document
from docx.shared import Inches

from visit.models import Visit, VisitContainer
from history.models import MedicalHistory
from datetime import datetime

class ReportGenerator(object):
    """"""
    styles = {}

    def __init__(self, response, visit_container):
        self.document = Document()
        self.buffer = response
        self.data = {
            'patient_name': visit_container.patient_detail.full_name,
            'patient_age': visit_container.patient_detail.age,
            'patient_sex': visit_container.patient_detail.sex,
            'visit_container': visit_container,
            'visits': visit_container.visits.all(),
        }

    def write_motor_tone(self, cns):
        self.document.add_heading('MOTOR TONE', level=3)
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'TONE'
        hdr_cells[1].text = 'Rt'
        hdr_cells[2].text = 'Lt'
        row_cells = table.add_row().cells
        row_cells[0].text = "UL"
        row_cells[1].text = cns.ms_tone_ul_r
        row_cells[2].text = cns.ms_tone_ul_l
        row_cells = table.add_row().cells
        row_cells[0].text = "LL"
        row_cells[1].text = cns.ms_tone_ll_r
        row_cells[2].text = cns.ms_tone_ll_r

    def write_motor_dtr(self,cns):
        self.document.add_heading('MOTOR DTR', level=3)
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'DTR'
        hdr_cells[1].text = 'Rt'
        hdr_cells[2].text = 'Lt'
        row_cells = table.add_row().cells
        row_cells[0].text = "B"
        row_cells[1].text = cns.ms_dtr_b_r
        row_cells[2].text = cns.ms_dtr_b_l
        row_cells = table.add_row().cells
        row_cells[0].text = "T"
        row_cells[1].text = cns.ms_dtr_t_r
        row_cells[2].text = cns.ms_dtr_t_l
        row_cells = table.add_row().cells
        row_cells[0].text = "S"
        row_cells[1].text = cns.ms_dtr_s_r
        row_cells[2].text = cns.ms_dtr_s_l
        row_cells = table.add_row().cells
        row_cells[0].text = "K"
        row_cells[1].text = cns.ms_dtr_k_r
        row_cells[2].text = cns.ms_dtr_k_l
        row_cells = table.add_row().cells
        row_cells[0].text = "A"
        row_cells[1].text = cns.ms_dtr_a_r
        row_cells[2].text = cns.ms_dtr_a_l
        row_cells = table.add_row().cells
        row_cells[0].text = "PLANTARS"
        row_cells[1].text = cns.ms_dtr_plorter_r
        row_cells[2].text = cns.ms_dtr_plorter_l

    def write_motor_power(self,cns):
        self.document.add_heading('MOTOR POWER', level=3)
        table = self.document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'POWER'
        hdr_cells[1].text = 'Rt'
        hdr_cells[2].text = 'Lt'
        row_cells = table.add_row().cells
        row_cells[0].text = "UL"
        row_cells[1].text = cns.ms_power_ul_r
        row_cells[2].text = cns.ms_power_ul_l
        row_cells = table.add_row().cells
        row_cells[0].text = "LL"
        row_cells[1].text = cns.ms_power_ll_r
        row_cells[2].text = cns.ms_power_ll_r

    def write_sensory(self,cns):
        self.document.add_heading('SENSORY SYSTEM', level=3)
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        row_cells = table.rows[0].cells
        row_cells[0].text = "TOUCH"
        row_cells[1].text = cns.ss_touch
        row_cells = table.add_row().cells
        row_cells[0].text = "TEMP"
        row_cells[1].text = cns.ss_temp
        row_cells = table.add_row().cells
        row_cells[0].text = "PAIN"
        row_cells[1].text = cns.ss_pain
        row_cells = table.add_row().cells
        row_cells[0].text = "GPS"
        row_cells[1].text = cns.ss_gps

        self.document.add_paragraph('Cerebelleim: ' + cns.cerebelleim)
        self.document.add_paragraph('Gait: ' + cns.gait)




    def create_document(self):
        self.document.add_heading('Patient Examination Report', 0)
        self.document.add_heading('Patient Information', level=1)
        # Patient Name is added first
        paragraph = self.document.add_paragraph()
        paragraph.add_run('Name').bold = True
        paragraph.add_run('  : ' + self.data['patient_name'])

        # Patient Age
        paragraph = self.document.add_paragraph()
        paragraph.add_run('Age').bold = True
        paragraph.add_run('   : ' + self.data['patient_age'])

        # Patient Sex
        paragraph = self.document.add_paragraph()
        paragraph.add_run('Sex').bold = True
        paragraph.add_run('   : ' + self.data['patient_sex'])

        # Report Date
        paragraph = self.document.add_paragraph()
        paragraph.add_run('Generated at').bold = True
        now = datetime.now()
        paragraph.add_run('  : ' + now.strftime("%m/%d/%Y, %H:%M:%S"))

        # Add diagnosis first
        visit = self.data['visits'][0]
        diagnose = visit.diagnose.all();
        self.document.add_heading('Diagnosis', level=2)
        for d in diagnose:
            final_str = d.name
            self.document.add_paragraph(final_str, style='List Bullet')

        # complaints
        self.document.add_heading('CHIEF Complaints', level=2)
        complaints = visit.complaints.all();
        for c in complaints:
            final_str = c.description + ' for ' + c.duration
            self.document.add_paragraph(final_str, style='List Bullet')

        self.document.add_heading('Medical History', level=1)
        history = MedicalHistory.objects.filter(
            patient_detail=self.data['visit_container'].patient_detail.patient_id)
        for h in history:
            self.document.add_paragraph(h.status, style='List Bullet')

        visits = self.data['visits']
        visit_count = 1;
        visit_size = len(visits)
        for visit in visits:
            self.document.add_page_break()
            self.document.add_heading('Visit On ' + visit.visit_date.strftime("%m/%d/%Y, %H:%M:%S"), level=1)
            self.document.add_heading('Examination at the time of admission', level=2);
            # Add diagnosis first
            # diagnose = visit.diagnose.all();
            # final_str = ", ".join([x.name for x in diagnose])
            # self.document.add_heading('Diagnosis', level=2)
            # self.document.add_paragraph(final_str)
            # weight      = visit.vitals.weight or ""
            # p_ce_cn_iol = visit.vitals.p_ce_cn_iol or "NAD"
            # temp        = visit.vitals.temp or "NAD"
            # pulse       = visit.vitals.pulse or "NAD"
            # bp          = visit.vitals.bp or "NAD"
            # rr          = visit.vitals.rr or "NAD"
            # # cns         = visit.vitals.cns or "NAD"
            # chest       = visit.vitals.chest or "NAD"
            # cvs         = visit.vitals.cvs or "NAD"
            # pa          = visit.vitals.pa or "NAD"
            # tests       = visit.vitals.tests or "NAD"

            self.document.add_heading('Vitals', level=2)
            table = self.document.add_table(rows=1, cols=2)
            row_cells = table.rows[0].cells
            row_cells[0].text = "Weight"
            row_cells[1].text = ""
            row_cells = table.add_row().cells
            row_cells[0].text = "P/CE/CN/IOL"
            row_cells[1].text = ""
            row_cells = table.add_row().cells
            row_cells[0].text = ""
            row_cells[1].text = ""
            row_cells = table.add_row().cells
            row_cells[0].text = "Pulse"
            row_cells[1].text = ""
            row_cells = table.add_row().cells
            row_cells[0].text = "BP - mmhg"
            row_cells[1].text = ""
            row_cells = table.add_row().cells
            row_cells[0].text = "RR /min"
            row_cells[1].text = ""
            row_cells = table.add_row().cells
            row_cells[0].text = "Chest b/l nvbs"
            row_cells[1].text = ""
            row_cells = table.add_row().cells
            row_cells[0].text = "CVS- s1 s2 N"
            row_cells[1].text = ""
            row_cells = table.add_row().cells
            row_cells[0].text = "PA- SOFT , NO ORGANOMEGALY"
            row_cells[1].text = ""

            self.document.add_heading('CNS', level=2)
            cns = visit.cns;
            if cns:
                self.document.add_paragraph('GCS : ' + (cns.gcs or ""))
                self.document.add_paragraph('HMF : ' + (cns.hmf or ""))
                self.document.add_paragraph('MOTOR SYSTEM')
                self.document.add_paragraph('MUSCLE BULK -  ' + (cns.muscle_bulk or "") )
                self.document.add_paragraph('NUTRITION -  ' + ( cns.nutrition or ""))
                self.write_motor_tone(cns)
                self.write_motor_dtr(cns)
                self.write_motor_power(cns)
                self.write_sensory(cns)

            # self.document.add_heading('Complaints', level=2)
            # complaints = visit.complaints.all();
            # for c in complaints:
            #     final_str = c.description + ' for ' + c.duration
            #     self.document.add_paragraph(final_str, style='List Bullet')

            self.document.add_heading('COURSE IN HOSPITAL STAY', level=2)
            self.document.add_heading('Medicine', level=2)
            medicines = visit.medicines.all();
            for m in medicines:
                self.document.add_paragraph(m.getDisplayValue(), style='List Bullet')

            self.document.add_heading('Remark', level=2)
            final_str = visit.remarks
            self.document.add_paragraph(final_str)

        self.document.save(self.buffer)


def generate_report(request, pk):
    visit_container = VisitContainer.objects.get(id=pk);
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="%s-%s.docx"' % (
        visit_container.patient_detail.full_name, visit_container.visit_date.strftime('%d-%m-%y'))
    doc = ReportGenerator(response, visit_container)
    doc.create_document()
    return response
